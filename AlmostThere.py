import docx
#  set the file path to the text file of the report.  run this.  and then run_a_macro to set the pages. 

with open(r"C:\Users\nsloss\Documents\certifieds\20230812", "r") as f:

        contents = f.readlines() #list containing all strings
        indexes = [contents.index(x) for x in set(contents)]
        projects =[]
        page_location = []
        wrd = 'PROJECT ID'
        wrd2='TOTALS'
        wrd3='Selection'
      
        for line in contents:
          if wrd3 in line:
                contents.remove(line)

      
        for line in contents:   
            if wrd in line:  
                if line not in projects:
                    if wrd2 in line:
                         pass
                    else:
                        line=line[15:22]
                        if line not in projects:
                            projects.append(line)
                
       
        for num in projects:
             for line in contents:
                  if num in line:
                       id = contents.index(line)
                       if id not in page_location:
                            page_location.append(id)
                    

                  

        print(projects)
        final_len =len(contents)+10
        print(final_len)
        page_location.append(final_len)
        print(page_location)
     
        loc = 0

        for num in projects:
          #with open( num +".txt", 'w') as file:
               loc1 = int(page_location[loc])-10
               loc2 = int(page_location[loc+1])-10
               loc+=1
               text=contents[loc1:loc2]
               
               doc = docx.Document("Template_for_Certifieds.docx")  
              
               p = doc.add_paragraph()
               p.paragraph_format.line_spacing = 1
               p.paragraph_format.space_after = 0           
               for line in text:
                         run = p.add_run(line)
                         run.font.name = 'Courier New'
                         run.font.size = docx.shared.Pt(7)   
                                         
               doc.save(num+".docx")
               
              
               #for line in contents[loc1:loc2]:
                # file.write(line)



"""
text=contents[0:66]

doc = docx.Document("Template_for_Certifieds.docx")
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0
for line in text:
     run = p.add_run(line)
     run.font.name = 'Courier New'
     run.font.size = docx.shared.Pt(7)
doc.save("helloSample.docx")

"""